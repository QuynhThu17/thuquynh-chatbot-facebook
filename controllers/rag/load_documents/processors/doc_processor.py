"""
DOC/DOCX Document Processor
Xử lý file Word (.doc, .docx), trích xuất text và ảnh theo đúng thứ tự
"""

import logging
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
import zipfile
from io import BytesIO
import xml.etree.ElementTree as ET
import re
import os
import sys
from pathlib import Path

# Add project directories to path
current_dir = Path(__file__).parent
load_documents_dir = current_dir.parent
sys.path.append(str(load_documents_dir))

from .base_processor import BaseDocumentProcessor, DocumentContent, PageContent
from controllers.rag.load_documents.utils import DocumentUtils

logger = logging.getLogger(__name__)

class DocProcessor(BaseDocumentProcessor):
    """
    Processor cho file DOC/DOCX
    Sử dụng python-docx để trích xuất text và ảnh
    """
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx', '.doc']
    
    def can_process(self, file_extension: str) -> bool:
        """Kiểm tra có thể xử lý file DOC/DOCX không"""
        return file_extension.lower() in self.supported_extensions
    
    async def extract_content(self, file_path: str, file_data: bytes) -> DocumentContent:
        """
        Trích xuất nội dung từ file DOC/DOCX
        
        Args:
            file_path: Đường dẫn file
            file_data: Dữ liệu file Word
            
        Returns:
            DocumentContent: Nội dung đã trích xuất
        """
        try:
            # Validate file
            if not self.validate_file(file_data, '.docx'):
                raise ValueError("Invalid Word document")
            
            # Xử lý DOCX
            if file_path.lower().endswith('.docx'):
                return await self._process_docx(file_path, file_data)
            else:
                # Cho DOC cũ, convert sang DOCX hoặc sử dụng python-docx2txt
                return await self._process_doc(file_path, file_data)
                
        except Exception as e:
            logger.error(f"Error extracting content from Word document: {str(e)}")
            raise
    
    async def _process_docx(self, file_path: str, file_data: bytes) -> DocumentContent:
        """
        Xử lý file DOCX
        
        Args:
            file_path: Đường dẫn file
            file_data: Dữ liệu file
            
        Returns:
            DocumentContent: Nội dung đã trích xuất
        """
        try:
            # Mở document
            doc = Document(BytesIO(file_data))
            
            # Trích xuất ảnh từ document relationships (nếu process_images=True)
            images = []
            if self.process_images:
                images = await self._extract_images_from_docx(file_data)
            
            # Trích xuất text và merge với ảnh
            text_content, merged_images = await self._extract_text_with_images(doc, images)
            
            # Tạo metadata
            metadata = self.extract_metadata(file_path, file_data)
            metadata.update(await self._extract_docx_metadata(doc))
            
            # Convert to new page-aware structure
            return self._convert_legacy_to_page_structure(text_content, merged_images, metadata)
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {str(e)}")
            raise
    
    async def _process_doc(self, file_path: str, file_data: bytes) -> DocumentContent:
        """
        Xử lý file DOC cũ (sử dụng fallback method)
        
        Args:
            file_path: Đường dẫn file
            file_data: Dữ liệu file
            
        Returns:
            DocumentContent: Nội dung đã trích xuất
        """
        try:
            # Sử dụng python-docx để thử đọc DOC
            # Note: python-docx chỉ hỗ trợ DOCX tốt, DOC có thể không đầy đủ
            doc = Document(BytesIO(file_data))
            
            # Trích xuất text đơn giản
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            text_content = self.clean_text(text_content)
            
            # Tạo metadata
            metadata = self.extract_metadata(file_path, file_data)
            metadata["note"] = "DOC file processed with limited support"
            
            # Convert to new page-aware structure
            return self._convert_legacy_to_page_structure(text_content, [], metadata)
            
        except Exception as e:
            logger.error(f"Error processing DOC: {str(e)}")
            # Fallback: chỉ trả về metadata
            metadata = {"error": f"Unable to process DOC file: {str(e)}"}
            return self._convert_legacy_to_page_structure("", [], metadata)
    
    async def _extract_images_from_docx(self, file_data: bytes) -> List[Dict[str, Any]]:
        """
        Trích xuất ảnh từ file DOCX
        
        Args:
            file_data: Dữ liệu file DOCX
            
        Returns:
            List[Dict]: Danh sách ảnh
        """
        images = []
        
        try:
            # Mở DOCX như ZIP file
            with zipfile.ZipFile(BytesIO(file_data), 'r') as zip_file:
                # Tìm tất cả file ảnh trong media folder
                media_files = [f for f in zip_file.namelist() if f.startswith('word/media/')]
                
                for i, media_file in enumerate(media_files):
                    try:
                        # Đọc dữ liệu ảnh
                        img_data = zip_file.read(media_file)
                        
                        # Validate ảnh
                        is_valid, img_metadata = DocumentUtils.validate_image_data(img_data)
                        
                        if is_valid:
                            # Tối ưu hóa ảnh
                            optimized_img_data = DocumentUtils.optimize_image(img_data)
                            
                            # Lấy extension từ filename
                            _, ext = os.path.splitext(media_file)
                            
                            images.append({
                                "image_data": optimized_img_data,
                                "position": i * 100,  # Estimate position
                                "page": 1,  # Word không có page concept rõ ràng
                                "image_index": i,
                                "metadata": {
                                    "original_filename": os.path.basename(media_file),
                                    "file_extension": ext,
                                    **img_metadata
                                }
                            })
                            
                    except Exception as e:
                        logger.warning(f"Error extracting image {media_file}: {str(e)}")
                        continue
                        
        except Exception as e:
            logger.error(f"Error extracting images from DOCX: {str(e)}")
        
        return images
    
    async def _extract_text_with_images(self, doc: Document, images: List[Dict[str, Any]]) -> tuple[str, List[Dict[str, Any]]]:
        """
        Trích xuất text và merge với ảnh theo thứ tự
        
        Args:
            doc: Document object
            images: Danh sách ảnh
            
        Returns:
            tuple: (text_content, merged_images)
        """
        text_content = ""
        image_counter = 0
        merged_images = []
        
        try:
            # Duyệt qua tất cả paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                
                if para_text:
                    text_content += para_text + "\n"
                
                # Kiểm tra có ảnh trong paragraph không
                if self._paragraph_has_image(paragraph):
                    if image_counter < len(images):
                        img_data = images[image_counter].copy()
                        img_data["position"] = len(text_content)
                        merged_images.append(img_data)
                        
                        # Chèn placeholder
                        text_content += f"[IMAGE_{image_counter + 1}]\n"
                        image_counter += 1
                
                text_content += "\n"
            
            # Thêm ảnh còn lại nếu có
            while image_counter < len(images):
                img_data = images[image_counter].copy()
                img_data["position"] = len(text_content)
                merged_images.append(img_data)
                text_content += f"[IMAGE_{image_counter + 1}]\n"
                image_counter += 1
            
            text_content = self.clean_text(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting text with images: {str(e)}")
            # Fallback: extract text only
            text_content = "\n".join([para.text for para in doc.paragraphs])
            text_content = self.clean_text(text_content)
            merged_images = images
        
        return text_content, merged_images
    
    def _paragraph_has_image(self, paragraph) -> bool:
        """
        Kiểm tra paragraph có chứa ảnh không
        
        Args:
            paragraph: Paragraph object
            
        Returns:
            bool: True nếu có ảnh
        """
        try:
            # Kiểm tra runs trong paragraph
            for run in paragraph.runs:
                if run.element.xpath('.//pic:pic'):
                    return True
                    
            # Kiểm tra drawing objects
            if paragraph.element.xpath('.//w:drawing'):
                return True
                
        except Exception:
            pass
            
        return False
    
    async def _extract_docx_metadata(self, doc: Document) -> Dict[str, Any]:
        """
        Trích xuất metadata từ DOCX
        
        Args:
            doc: Document object
            
        Returns:
            Dict[str, Any]: Metadata
        """
        try:
            core_props = doc.core_properties
            
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "category": core_props.category or "",
                "comments": core_props.comments or "",
                "created": core_props.created.isoformat() if core_props.created else "",
                "modified": core_props.modified.isoformat() if core_props.modified else "",
                "last_modified_by": core_props.last_modified_by or "",
                "revision": core_props.revision or 0,
                "paragraph_count": len(doc.paragraphs),
                "has_tables": len(doc.tables) > 0,
                "table_count": len(doc.tables)
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {str(e)}")
            return {}
    
    async def _extract_table_content(self, doc: Document) -> str:
        """
        Trích xuất nội dung từ tables trong document
        
        Args:
            doc: Document object
            
        Returns:
            str: Table content
        """
        table_content = ""
        
        try:
            for table in doc.tables:
                table_content += "\n[TABLE]\n"
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_data.append(cell_text)
                    
                    if row_data:
                        table_content += " | ".join(row_data) + "\n"
                
                table_content += "[/TABLE]\n"
                
        except Exception as e:
            logger.error(f"Error extracting table content: {str(e)}")
        
        return table_content
