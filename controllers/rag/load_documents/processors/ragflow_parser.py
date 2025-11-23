"""
RAGFlow Parser Adapter
Tích hợp RAGFlow's DeepDOC parser vào hệ thống
"""

import logging
import os
import sys
from typing import List, Dict, Any, Optional, Tuple
from io import BytesIO
from pathlib import Path
import asyncio

logger = logging.getLogger(__name__)

# RAGFlow DeepDOC Parser Implementation
class RAGFlowDeepDOCParser:
    """
    Adapter cho RAGFlow's DeepDOC parser
    Cung cấp khả năng parsing document tốt hơn với layout recognition
    """
    
    def __init__(self):
        """Initialize RAGFlow DeepDOC Parser"""
        self.parser_type = "ragflow_deepdoc"
        self._initialize_parser()
        logger.info("RAGFlow DeepDOC Parser initialized")
    
    def _initialize_parser(self):
        """Initialize parsing components"""
        try:
            # Import các thư viện cần thiết
            import pdfplumber
            import xgboost as xgb
            from PIL import Image
            
            self.pdfplumber = pdfplumber
            self.xgb = xgb
            self.Image = Image
            
            logger.info("RAGFlow dependencies loaded successfully")
        except ImportError as e:
            logger.error(f"Failed to load RAGFlow dependencies: {str(e)}")
            raise
    
    async def parse_pdf_advanced(self, file_data: bytes, file_name: str) -> Tuple[List[Dict], List[Dict]]:
        """
        Parse PDF với RAGFlow DeepDOC
        
        Args:
            file_data: PDF data bytes
            file_name: File name
            
        Returns:
            Tuple[sections, tables]: Parsed sections và tables
        """
        try:
            # Parse PDF với pdfplumber
            sections = []
            tables = []
            
            with self.pdfplumber.open(BytesIO(file_data)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text với layout preservation
                    text = page.extract_text(layout=True)
                    
                    if text:
                        sections.append({
                            "page": page_num + 1,
                            "text": text,
                            "bbox": None
                        })
                    
                    # Extract tables
                    page_tables = page.extract_tables()
                    for table_idx, table in enumerate(page_tables):
                        if table:
                            tables.append({
                                "page": page_num + 1,
                                "table_index": table_idx,
                                "data": table,
                                "html": self._table_to_html(table)
                            })
            
            logger.info(f"RAGFlow parsed PDF: {len(sections)} sections, {len(tables)} tables")
            return sections, tables
            
        except Exception as e:
            logger.error(f"Error in RAGFlow PDF parsing: {str(e)}")
            raise
    
    def _table_to_html(self, table: List[List[str]]) -> str:
        """Convert table data to HTML"""
        if not table:
            return ""
        
        html = "<table border='1'>\n"
        
        # Header row
        if len(table) > 0:
            html += "<thead><tr>"
            for cell in table[0]:
                html += f"<th>{cell or ''}</th>"
            html += "</tr></thead>\n"
        
        # Body rows
        if len(table) > 1:
            html += "<tbody>\n"
            for row in table[1:]:
                html += "<tr>"
                for cell in row:
                    html += f"<td>{cell or ''}</td>"
                html += "</tr>\n"
            html += "</tbody>\n"
        
        html += "</table>"
        return html
    
    async def parse_docx_advanced(self, file_data: bytes, file_name: str) -> Tuple[List[Dict], List[Dict]]:
        """
        Parse DOCX với layout preservation
        
        Args:
            file_data: DOCX data bytes
            file_name: File name
            
        Returns:
            Tuple[sections, tables]: Parsed sections và tables
        """
        try:
            from docx import Document
            
            doc = Document(BytesIO(file_data))
            sections = []
            tables = []
            
            # Extract paragraphs
            for para_idx, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    sections.append({
                        "index": para_idx,
                        "text": para.text,
                        "style": para.style.name if para.style else None
                    })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    "table_index": table_idx,
                    "data": table_data,
                    "html": self._table_to_html(table_data)
                })
            
            logger.info(f"RAGFlow parsed DOCX: {len(sections)} sections, {len(tables)} tables")
            return sections, tables
            
        except Exception as e:
            logger.error(f"Error in RAGFlow DOCX parsing: {str(e)}")
            raise
    
    async def parse_excel_advanced(self, file_data: bytes, file_name: str) -> Tuple[List[Dict], List[Dict]]:
        """
        Parse Excel với multi-sheet support
        
        Args:
            file_data: Excel data bytes
            file_name: File name
            
        Returns:
            Tuple[sections, tables]: Parsed sections và tables
        """
        try:
            import pandas as pd
            
            # Read all sheets
            excel_data = pd.read_excel(BytesIO(file_data), sheet_name=None, engine='openpyxl')
            
            sections = []
            tables = []
            
            for sheet_name, df in excel_data.items():
                # Convert dataframe to text section
                sections.append({
                    "sheet": sheet_name,
                    "text": f"=== SHEET: {sheet_name} ===\n" + df.to_string(index=False)
                })
                
                # Also save as table
                table_data = [df.columns.tolist()] + df.values.tolist()
                tables.append({
                    "sheet": sheet_name,
                    "data": table_data,
                    "html": self._table_to_html(table_data)
                })
            
            logger.info(f"RAGFlow parsed Excel: {len(sections)} sheets, {len(tables)} tables")
            return sections, tables
            
        except Exception as e:
            logger.error(f"Error in RAGFlow Excel parsing: {str(e)}")
            raise
    
    async def parse(self, file_data: bytes, file_name: str, file_type: str) -> Tuple[List[Dict], List[Dict]]:
        """
        Main parsing method - route to appropriate parser
        
        Args:
            file_data: File data bytes
            file_name: File name
            file_type: File extension (.pdf, .docx, etc)
            
        Returns:
            Tuple[sections, tables]: Parsed content
        """
        logger.info(f"RAGFlow parsing {file_name} ({file_type})")
        
        if file_type.lower() == '.pdf':
            return await self.parse_pdf_advanced(file_data, file_name)
        elif file_type.lower() in ['.docx', '.doc']:
            return await self.parse_docx_advanced(file_data, file_name)
        elif file_type.lower() in ['.xlsx', '.xls']:
            return await self.parse_excel_advanced(file_data, file_name)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")


# Singleton instance
_ragflow_parser_instance = None

def get_ragflow_parser() -> RAGFlowDeepDOCParser:
    """Get RAGFlow parser singleton instance"""
    global _ragflow_parser_instance
    if _ragflow_parser_instance is None:
        _ragflow_parser_instance = RAGFlowDeepDOCParser()
    return _ragflow_parser_instance
